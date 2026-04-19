import os
import re
import sys
from pathlib import Path
from datetime import datetime
from difflib import SequenceMatcher

# Add project root to path if running directly
sys.path.append(str(Path(__file__).parent.parent))

import PyPDF2
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from fpdf import FPDF
import shutil

import config
import yaml
from src.llm_provider import get_llm

def _extract_content_from_tags(text: str, tag_name: str) -> str:
    """Extract content between <tag> and </tag>."""
    pattern = rf"<{tag_name}>(.*?)</{tag_name}>"
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return text.strip()


def parse_resume(file_path: Path = None) -> str:
    """
    Read the base resume from PDF or DOCX and return as plain text.
    Auto-detects format based on file extension.
    """
    if file_path is None:
        # Auto-detect: prefer DOCX, fall back to PDF
        if config.BASE_RESUME_DOCX.exists():
            file_path = config.BASE_RESUME_DOCX
        elif config.BASE_RESUME_PDF.exists():
            file_path = config.BASE_RESUME_PDF
        else:
            raise FileNotFoundError(
                "No base resume found. Place your resume as "
                "'data/base_resume.pdf' or 'data/base_resume.docx'"
            )

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(file_path)
    else:
        # Try reading as plain text
        return file_path.read_text(encoding="utf-8")


def _parse_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    text_parts = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _parse_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    doc = Document(path)
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def _load_prompt_template(template_name: str) -> str:
    """Load a prompt template from the templates directory."""
    template_path = config.TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    return template_path.read_text(encoding="utf-8")


def tailor_resume(
    base_resume_text: str,
    job_title: str,
    company: str,
    location: str,
    job_description: str,
    profile_data: dict = None
) -> str:
    """Use the LLM to tailor the resume, injecting real profile data."""
    template = _load_prompt_template("resume_prompt.txt")
    
    # Pre-inject profile data to remove placeholders
    full_name = f"{profile_data.get('personal', {}).get('first_name', '')} {profile_data.get('personal', {}).get('last_name', '')}"
    email = profile_data.get('personal', {}).get('email', '')
    phone = profile_data.get('personal', {}).get('phone', '')
    linked_in = profile_data.get('personal', {}).get('linkedin_url', '')
    p_location = f"{profile_data.get('personal', {}).get('city', '')}, {profile_data.get('personal', {}).get('country', '')}"

    # First, format the template with the core job data
    prompt = template.format(
        base_resume=base_resume_text,
        job_title=job_title,
        company=company,
        location=location,
        job_description=job_description,
    )

    # Phase 24.1: ROBUST PLACEHOLDER INJECTION
    p = prompt.replace("{{full_name}}", full_name)
    p = p.replace("{{email}}", email)
    p = p.replace("{{phone}}", phone)
    p = p.replace("{{linkedin_url}}", linked_in)
    p = p.replace("{{location}}", p_location)
    
    # Handle single-brace versions just in case LLM or formatting stripped them
    p = p.replace("{full_name}", full_name)
    p = p.replace("{email}", email)
    p = p.replace("{phone}", phone)
    p = p.replace("{linkedin_url}", linked_in)
    p = p.replace("{location}", p_location)
    prompt = p


    llm = get_llm()
    print(f"  ⚙ Performing Surgical Experience Surgery with {config.LLM_PROVIDER}...")

    # The Surgical Prompt: Specifically refining experience bullet points
    system_msg = (
        "You are an Elite Career Surgeon. Your job is to rewrite the EXPERIENCE section of the resume "
        "to perfectly match the Job Description. Use the exact keywords 'vibe' and terminology of the JD. "
        "Highlight achievements that matter to THIS employer. Output ONLY the tailored resume in markdown "
        "inside <resume> tags. Keep all personal info and contact details as provided."
    )

    result = llm.generate(prompt, system_prompt=system_msg)
    return _extract_content_from_tags(result, "resume")


def generate_cover_letter(
    base_resume_text: str,
    job_title: str,
    company: str,
    location: str,
    job_description: str,
    profile_data: dict = None
) -> str:
    """Generate a tailored cover letter with current date and name."""
    template = _load_prompt_template("cover_letter_prompt.txt")
    
    current_date = datetime.now().strftime("%B %d, %Y")
    first_name = profile_data.get('personal', {}).get('first_name', '')
    last_name = profile_data.get('personal', {}).get('last_name', '')
    full_name = f"{first_name} {last_name}"
    email = profile_data.get('personal', {}).get('email', '')
    phone = profile_data.get('personal', {}).get('phone', '')
    location_str = f"{profile_data.get('personal', {}).get('city', '')}, {profile_data.get('personal', {}).get('province', '')}"
    linkedin = profile_data.get('personal', {}).get('linkedin_url', '')

    prompt = template.format(
        base_resume=base_resume_text,
        job_title=job_title,
        company=company,
        location=location,
        job_description=job_description,
    )

    llm = get_llm()
    print(f"  ⚙ Generating cover letter with {config.LLM_PROVIDER}...")
    result = llm.generate(
        prompt,
        system_prompt="You are an expert executive speechwriter. Write the BODY of a persuasive, professional cover letter inside <cover_letter> tags. Do NOT include a header or signature, as they will be added automatically."
    )
    body = _extract_content_from_tags(result, "cover_letter")

    # Phase 24.1: Manual Header & Signature Injection (Guaranteed Accuracy)
    header = f"# {full_name}\n{email} | {phone} | {location_str} | {linkedin}\n\n"
    cl_content = f"{header}{current_date}\n\nDear Hiring Manager,\n\n{body}\n\nSincerely,\n\n{full_name}"
    return cl_content



def generate_interview_prep(
    base_resume_text: str,
    job_title: str,
    company: str,
    location: str,
    job_description: str,
) -> str:
    """
    Use the LLM to generate an interview preparation cheat sheet.
    Returns the cheat sheet as markdown text.
    """
    template = _load_prompt_template("interview_prep_prompt.txt")
    prompt = template.format(
        base_resume=base_resume_text,
        job_title=job_title,
        company=company,
        location=location,
        job_description=job_description,
    )

    llm = get_llm()
    print(f"  ⚙ Generating interview prep with {config.LLM_PROVIDER}...")
    result = llm.generate(
        prompt,
        system_prompt="You are a senior technical recruiter. Output a detailed interview prep guide in markdown."
    )
    return result


def _sanitize_filename(name: str) -> str:
    """Make a string safe for use in a filename."""
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    name = re.sub(r'\s+', '_', name.strip())
    return name[:60]  # Limit length


def _markdown_to_docx(markdown_text: str, output_path: Path, title: str = ""):
    """Convert markdown text to a formatted DOCX document with a professional header."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri" # More modern ATS font
    font.size = Pt(10.5)
    
    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = markdown_text.split("\n")
    header_processed = False

    for line in lines:
        stripped = line.strip()
        
        # Legacy Mirror Header Logic (Robust Left-Alignment)
        if stripped.startswith("# ") and not header_processed:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(22)
            header_processed = 1
            continue

        if header_processed == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(stripped).font.size = Pt(10)
            header_processed = 2
            continue
            
        if header_processed == 2:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(stripped).font.size = Pt(10)
            
            # Draw Divider and finish header
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            p.add_run("_" * 100).font.size = Pt(6)
            
            header_processed = "DONE"
            continue

        if not stripped:
            doc.add_paragraph("")
            continue
            
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            # Spacing before sections
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(stripped[3:].upper())
            run.bold = True
            run.font.size = Pt(12)
        elif stripped.startswith("### "):
            # Phase 24.1: Unified Dual-Column Experience
            parts = [p.strip() for p in stripped[4:].split("|")]
            p = doc.add_paragraph()
            if len(parts) >= 2:
                # Use tab stops for right alignment effect in Word
                p.paragraph_format.tab_stops.add_tab_stop(Inches(7.25), WD_TAB_ALIGNMENT.RIGHT)
                run = p.add_run(parts[0])
                run.bold = True
                p.add_run("\t")
                p.add_run(" | ".join(parts[1:])).bold = True
            else:
                p.add_run(stripped[4:]).bold = True
        elif stripped.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            # Clear residual markdown from AI
            clean_line = stripped[2:].replace("#", "").replace("**", "").replace("__", "").strip()
            p.add_run(clean_line)
        else:
            doc.add_paragraph(stripped)


        # Regular Markdown Parsing
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(12)
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(11)
        elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.2)
            # Support inline bold in bullets
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        elif stripped.startswith("---"):
            p = doc.add_paragraph()
            p.add_run("_" * 75)
        else:
            p = doc.add_paragraph()
            # Inline bold: **text**
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(str(output_path))


def _sanitize_for_pdf(text: str) -> str:
    """Remove characters that FPDF's built-in fonts can't render and strip AI-hallucinated markers."""
    # Phase 24.1: AGGRESSIVE MARKER PURGE (Regex)
    # Strip leading hashtags, asterisks, underscores, and dashes that are not bullet-intended
    text = re.sub(r'^[#\*_\-\s]+', '', text)
    # Strip any remaining bold/italic markers
    text = text.replace("**", "").replace("__", "").replace("`", "")
    
    # Replace common unicode with ASCII or Latin-1 equivalents
    replacements = {
        '\u2018': "'", '\u2019': "'",  # smart quotes
        '\u201c': '"', '\u201d': '"',
        '\u2013': '-', '\u2014': '--',  # dashes
        '\u2026': '...', 
        '\u2022': ' ', # Bullet will be drawn manually, so strip the character
        '\u00a0': ' ',  # non-breaking space
        '\u2023': '>', '\u25cf': ' ',
        '\u2794': '->', '\u2192': '->',
        '\u25aa': '-', '\u25ab': '-', # square bullets
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Strip any remaining non-latin-1 characters
    try:
        return text.encode('latin-1', errors='replace').decode('latin-1').strip()
    except Exception:
        return text.encode('ascii', errors='replace').decode('ascii').strip()


def _text_to_pdf(text: str, output_path: Path):

    """Convert markdown text to a professional PDF document with robust styling."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin

    lines = text.split("\n")
    header_processed = False
    last_header_text = ""

    for line in lines:
        # Phase 24.1: CRITICAL MARGIN RESET
        pdf.set_x(pdf.l_margin)
        
        stripped = line.strip()
        if not stripped:
            pdf.ln(2)
            continue

        # Legacy Mirror Header Logic (Left-Aligned Name)
        if stripped.startswith("# ") and not header_processed:
            pdf.set_font("Helvetica", "B", 22)
            pdf.set_text_color(20, 20, 20)
            pdf.multi_cell(usable_w, 12, _sanitize_for_pdf(stripped[2:]), align="L")
            header_processed = 1 # Line 1 of contact info
            continue
        
        if header_processed == 1:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(usable_w, 5, _sanitize_for_pdf(stripped), align="L")
            header_processed = 2 # Line 2
            continue
            
        if header_processed == 2:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(usable_w, 5, _sanitize_for_pdf(stripped), align="L")
            pdf.ln(1)
            # Full width divider
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(4)
            header_processed = "DONE"
            continue

        pdf.set_text_color(0, 0, 0)
        if stripped.startswith("## "):
            # Phase 25.0: Section Header Buffer (40mm)
            if pdf.get_y() > (pdf.h - 40):
                pdf.add_page()
                
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(20, 20, 20)
            pdf.multi_cell(usable_w, 8, _sanitize_for_pdf(stripped[3:]).upper())
            pdf.set_font("Helvetica", "", 10)
            
        elif stripped.startswith("### "):
            # Phase 25.0: Job Header Buffer (50mm - Header + 2 Bullets)
            if pdf.get_y() > (pdf.h - 50):
                pdf.add_page()
                
            # Phase 24.1: Dual-Column Experience Header

            parts = [p.strip() for p in stripped[4:].split("|")]
            pdf.set_font("Helvetica", "B", 10.5)
            last_header_text = parts[0] if parts else stripped[4:]
            if len(parts) >= 2:
                title = parts[0]
                rest = " | ".join(parts[1:])
                pdf.cell(usable_w * 0.6, 6, _sanitize_for_pdf(title), ln=0)
                pdf.cell(usable_w * 0.4, 6, _sanitize_for_pdf(rest), ln=1, align="R")
            else:
                pdf.multi_cell(usable_w, 6, _sanitize_for_pdf(stripped[4:]))
            pdf.set_font("Helvetica", "", 10)

        elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
            content = _sanitize_for_pdf(stripped[2:])
            
            # Phase 24.1: SMART DUPLICATE FILTER
            # Skip if the bullet content near-matches the job title (common AI duplication)
            if last_header_text and len(content) > 0:
                similarity = SequenceMatcher(None, content.lower(), last_header_text.lower()).ratio()
                if similarity > 0.8 or last_header_text.lower() in content.lower():
                    continue

            pdf.set_font("Helvetica", "", 10)
            bullet_indent = 7 # Increased for executive look
            
            # Phase 24.1: ELITE BOLDER BULLETS
            orig_l_margin = pdf.l_margin
            pdf.set_left_margin(orig_l_margin + bullet_indent)
            pdf.set_x(orig_l_margin + bullet_indent)
            
            # Larger, bolder circle (1.8mm diameter)
            # Centering: font height at 10pt is ~3.5mm, so (3.5 - 1.8)/2 ~= 0.85 offset
            bullet_y = pdf.get_y() + 1.5
            pdf.set_fill_color(0, 0, 0) # True black for impact
            pdf.ellipse(orig_l_margin + 2.0, bullet_y, 1.8, 1.8, 'F')
            
            # Draw content with hanging indent
            pdf.multi_cell(usable_w - bullet_indent, 5, content)
            
            # Reset margin
            pdf.set_left_margin(orig_l_margin)
            pdf.ln(0.5)


        elif stripped == "---":
            pdf.ln(2)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(2)
            
        else:
            pdf.set_font("Helvetica", "", 10.5)
            # Phase 24.1: Human-Grade Spacing & Alignment
            pdf.multi_cell(usable_w, 6.5, _sanitize_for_pdf(stripped), align="L")
            pdf.ln(1.5)

    pdf.output(str(output_path))




def generate_documents(
    job_title: str,
    company: str,
    location: str,
    job_description: str,
    base_resume_path: Path = None,
) -> dict:
    """
    Full pipeline: parse resume → tailor → generate cover letter → save files.
    
    Returns dict with paths to generated files:
    {
        "output_dir": Path,
        "resume_docx": Path,
        "resume_pdf": Path,
        "resume_pdf": Path,
        "cover_letter_docx": Path,
        "cover_letter_pdf": Path,
        "interview_prep_md": Path,
    }
    """
    # Parse base resume
    print("📄 Parsing base resume...")
    base_text = parse_resume(base_resume_path)

    # Create output directory organized by date
    today = datetime.now().strftime("%Y-%m-%d")
    date_dir = config.OUTPUT_DIR / today
    date_dir.mkdir(parents=True, exist_ok=True)
    
    timestamp = datetime.now().strftime("%H%M%S")
    safe_company = _sanitize_filename(company)
    safe_job = _sanitize_filename(job_title)
    
    output_dir = date_dir / f"{safe_company}_{safe_job}_{timestamp}"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Load profile data for data injection
    profile_data = {}
    if config.PROFILE_PATH.exists():
        with open(config.PROFILE_PATH, "r", encoding="utf-8") as f:
            profile_data = yaml.safe_load(f) or {}

    # Tailor resume
    print("✏️  Tailoring resume...")
    tailored_resume = tailor_resume(base_text, job_title, company, location, job_description, profile_data)

    # Generate cover letter
    print("📝 Generating cover letter...")
    cover_letter = generate_cover_letter(base_text, job_title, company, location, job_description, profile_data)

    # Generate interview prep
    print("💡 Generating interview prep guide...")
    interview_prep = generate_interview_prep(base_text, job_title, company, location, job_description)

    # Phase 32.8: Elite Professional Branding - Standardized Naming
    first_name = profile_data.get('personal', {}).get('first_name', 'Candidate')
    last_name = profile_data.get('personal', {}).get('last_name', '')
    safe_name = _sanitize_filename(f"{first_name} {last_name}").strip()
    
    # Clean the job title for the filename
    clean_job = job_title.split("(")[0].strip() # Remove extra info like (Remote) or (Contract)
    safe_role = _sanitize_filename(clean_job)
    if len(safe_role) > 25: safe_role = safe_role[:22] + "..."
    
    date_str = datetime.now().strftime("%Y-%m-%d")
    resume_name = f"{date_str} - {safe_name} - {safe_role}"
    cl_name = f"Cover Letter - {date_str} - {safe_name} - {safe_role}"

    # Target: Official Mission Output Hub (AppData)
    final_dir = config.OUTPUT_DIR / date_str / f"{_sanitize_filename(company)}_{_sanitize_filename(job_title)}"
    final_dir.mkdir(parents=True, exist_ok=True)

    paths = {
        "output_dir": final_dir,
        "resume_pdf": final_dir / f"{resume_name}.pdf",
        "resume_docx": final_dir / f"{resume_name}.docx",
        "cover_letter_pdf": final_dir / f"{cl_name}.pdf",
        "cover_letter_docx": final_dir / f"{cl_name}.docx",
        "interview_prep": final_dir / f"Interview Prep - {safe_name}.md"
    }

    # Internal Raw Cache
    (final_dir / "mission_intel_raw.md").write_text(f"# {job_title}\n\n## Bio\n{tailored_resume}\n\n## Outreach\n{cover_letter}", encoding="utf-8")
    
    # Save Interview Prep
    paths["interview_prep"].write_text(interview_prep, encoding="utf-8")

    # Generate DOCX Versions (The Master Format)
    _markdown_to_docx(tailored_resume, paths["resume_docx"], title=resume_name)
    _markdown_to_docx(cover_letter, paths["cover_letter_docx"], title=cl_name)

    # Save PDF Versions
    try:
        _text_to_pdf(tailored_resume, paths["resume_pdf"])
        _text_to_pdf(cover_letter, paths["cover_letter_pdf"])
    except Exception as e:
        print(f"  ⚠️  PDF Generation Warning: {e}")

    print(f"  ✓ Mission assets deployed to: {final_dir}")
    return paths


if __name__ == "__main__":
    config.validate()
    # Quick test with a sample job
    sample_job = {
        "title": "Senior Python Developer",
        "company": "Tech Corp",
        "location": "Remote",
        "description": "We are looking for a Senior Python Developer with 5+ years of experience..."
    }
    paths = generate_documents(
        job_title=sample_job["title"],
        company=sample_job["company"],
        location=sample_job["location"],
        job_description=sample_job["description"],
    )
    print("\nGenerated files:")
    for key, path in paths.items():
        print(f"  {key}: {path}")
